# =============================================================================
#  utils.py - Logique métier pour l'application MINFI Recettes Fiscales
# =============================================================================
import pandas as pd
import numpy as np
import joblib
import io
from pathlib import Path
from datetime import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

DATA_DIR     = Path(__file__).parent / "data"
EXCEL_DIR    = DATA_DIR / "excel"
MODELS_DIR   = DATA_DIR / "modeles"
FIGURES_DIR  = DATA_DIR / "figures"

# Palette synchronisée avec app.py
NAVY        = "#0F1535"
PURPLE      = "#4318FF"
PURPLE_LT   = "#7551FF"
ACCENT_GOLD = "#FFB547"
ACCENT_GREEN= "#01B574"
ACCENT_PINK = "#FF5757"
ACCENT_CYAN = "#39B8FF"
TEXT_DARK   = "#1B2559"
TEXT_GRAY   = "#68769F"

MOIS_MAP = {
    "Janvier": 1, "Fevrier": 2, "Février": 2, "Mars": 3, "Avril": 4,
    "Mai": 5, "Juin": 6, "Juillet": 7, "Aout": 8, "Août": 8,
    "Septembre": 9, "Octobre": 10, "Novembre": 11, "Decembre": 12, "Décembre": 12,
}

RENAME_KEEP = {
    "Années":                                                                 "annee",
    "Mois":                                                                   "mois",
    "Recettes fiscales":                                                      "Recettes_fiscales",
    "Temp_moy":                                                               "Temp_moy",
    "Depenses pub":                                                           "Depenses_pub",
    "Solde budg.":                                                            "Solde_budg",
    "Prix pétrole brut (Brent)prices in nominalUS dollar ($/bbl)":            "Prix_petrole_Brent",
    "Indice de Taux de change effectif réel (REER)Source:FMI":                "REER",
    "indice de prix alimentaire mondial(base 2014-2016)FAO":                  "Indice_prix_alim_FAO",
    "Prix mondial riz-composite 5% (thai-0,80 ,viet-0,20) ($/mt)":            "Prix_riz_composite",
    "indice des prix des engrais (based on nominal US dollars, 2010=100)":    "Indice_engrais",
    "Prix mondial du cacao ( $/kg)":                                          "Prix_cacao",
    "Coffee, Robusta($/kg)":                                                  "Prix_cafe_Robusta",
    "Precipitations":                                                         "Precipitations",
    "Cotton, A Index($/kg)":                                                  "Prix_coton",
}

# =============================================================================
# 1. CHARGEMENT DE DONNÉES
# =============================================================================
def load_default_data() -> pd.DataFrame:
    file_path = EXCEL_DIR / "Base_de_Travail_modelisation.xlsx"
    df_raw = pd.read_excel(file_path, sheet_name="données_mensuelles", header=1)

    keep = [c for c in RENAME_KEEP if c in df_raw.columns]
    df = df_raw[keep].rename(columns=RENAME_KEEP).copy()
    df["annee"] = df["annee"].ffill().astype(int)
    df["mois_num"] = df["mois"].astype(str).str.strip().map(MOIS_MAP)
    df.index = pd.PeriodIndex(
        [f"{int(r.annee)}-{int(r.mois_num):02d}"
         for _, r in df.iterrows() if pd.notna(r.mois_num)],
        freq="M",
    )
    df = df.drop(columns=["annee", "mois", "mois_num"])
    df = df.apply(pd.to_numeric, errors="coerce")
    df = df.interpolate(method="linear", limit_direction="both").ffill().bfill()
    return df


def parse_uploaded_file(uploaded_file) -> pd.DataFrame:
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    else:
        try:
            df_raw = pd.read_excel(uploaded_file, sheet_name="données_mensuelles", header=1)
            keep = [c for c in RENAME_KEEP if c in df_raw.columns]
            if "Recettes fiscales" in keep:
                df = df_raw[keep].rename(columns=RENAME_KEEP).copy()
                df["annee"] = df["annee"].ffill().astype(int)
                df["mois_num"] = df["mois"].astype(str).str.strip().map(MOIS_MAP)
                df.index = pd.PeriodIndex(
                    [f"{int(r.annee)}-{int(r.mois_num):02d}"
                     for _, r in df.iterrows() if pd.notna(r.mois_num)],
                    freq="M",
                )
                df = df.drop(columns=["annee", "mois", "mois_num"])
                df = df.apply(pd.to_numeric, errors="coerce")
                df = df.interpolate(method="linear", limit_direction="both").ffill().bfill()
                return df
        except Exception:
            pass
        df = pd.read_excel(uploaded_file)

    date_col = None
    for c in df.columns:
        cl = str(c).lower()
        if cl in ("date", "mois", "month", "période", "periode", "year_month"):
            date_col = c
            break

    if date_col is not None:
        df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
        df = df.dropna(subset=[date_col]).set_index(date_col)
        df.index = df.index.to_period("M")

    df = df.apply(pd.to_numeric, errors="coerce")
    df = df.interpolate(method="linear", limit_direction="both").ffill().bfill()
    return df


# =============================================================================
# 2. MODÈLES & METRICS
# =============================================================================
def load_metrics(horizon: int):
    file = EXCEL_DIR / "rapport_complet.xlsx"
    if not file.exists():
        return None
    try:
        return pd.read_excel(file, sheet_name=f"Metriques_H{horizon}")
    except Exception:
        return None


def load_previsions_2025(horizon: int):
    mapping = {
        3:  "previsions_2025_H3_Q1_XGBoost.xlsx",
        6:  "previsions_2025_H6_S1_XGBoost.xlsx",
        12: "previsions_2025_H12_Année complète_RandomForest.xlsx",
    }
    file = EXCEL_DIR / mapping[horizon]
    if not file.exists():
        return None
    return pd.read_excel(file)


def load_backtest(horizon: int, year: int = 2024):
    mapping = {
        3:  f"backtest_{year}_H3_Q1_XGBoost.xlsx",
        6:  f"backtest_{year}_H6_S1_XGBoost.xlsx",
        12: f"backtest_{year}_H12_Année complète_RandomForest.xlsx",
    }
    file = EXCEL_DIR / mapping[horizon]
    if not file.exists():
        return None
    return pd.read_excel(file)


# =============================================================================
# 3. PRÉVISION MULTI-ANNÉE (1, 2, 3 ans)
# =============================================================================
def extend_forecast_multi_year(df_hist: pd.DataFrame, n_years: int = 1):
    """
    Étend la prévision annuelle XGBoost (2025) à 2 ou 3 années en utilisant :
      - la base XGBoost 2025 (chargée depuis l'Excel)
      - une extrapolation linéaire du trend + profil saisonnier sur 2026 et 2027
      - élargissement progressif de l'intervalle de confiance par horizon
    Renvoie (prev_df, dict cumul par année).
    """
    base = load_previsions_2025(12)
    if base is None:
        return None, None

    base = base.copy()
    base["Mois"] = base["Mois"].astype(str)
    by_year = {2025: base["Prévision (Mds)"].sum()}

    if n_years == 1:
        return base, by_year

    # Préparation : tendance + saisonnalité depuis l'historique
    serie = df_hist["Recettes_fiscales"].dropna()
    if not isinstance(serie.index, pd.DatetimeIndex):
        serie.index = pd.to_datetime(serie.index)

    # Tendance linéaire (régression simple)
    x = np.arange(len(serie))
    slope, intercept = np.polyfit(x, serie.values, 1)
    trend_last = intercept + slope * (len(serie) - 1)

    # Profil saisonnier multiplicatif (mois -> coefficient autour de 1)
    monthly_mean = serie.groupby(serie.index.month).mean()
    seasonal_profile = monthly_mean / monthly_mean.mean()

    # Volatilité résiduelle pour les IC
    residual_std = (serie - (intercept + slope * x)).std()

    additional_years = list(range(2026, 2025 + n_years))
    extra_rows = []
    last_total_2025 = base["Prévision (Mds)"].sum()

    for k_year, year in enumerate(additional_years, start=1):
        year_total = 0.0
        for m in range(1, 13):
            # nombre de mois après la fin de l'historique
            h_offset = (year - serie.index[-1].year) * 12 + (m - serie.index[-1].month)
            level = trend_last + slope * h_offset
            value = level * seasonal_profile.get(m, 1.0)

            # IC s'élargit avec l'horizon (sqrt-temps)
            ic_half = 1.96 * residual_std * np.sqrt(1 + h_offset / 12.0)
            extra_rows.append({
                "Mois": f"{year}-{m:02d}",
                "Prévision (Mds)": round(value, 2),
                "IC bas 95%":      round(value - ic_half, 2),
                "IC haut 95%":     round(value + ic_half, 2),
                "Total cumulé":    None,  # rempli après
            })
            year_total += value
        by_year[year] = round(year_total, 2)

    extra_df = pd.DataFrame(extra_rows)
    full = pd.concat([base, extra_df], ignore_index=True)

    # Recalculer Total cumulé PAR ANNÉE (pas inter-années)
    cum_by_year = {}
    new_totals = []
    for _, row in full.iterrows():
        yr = int(row["Mois"].split("-")[0])
        cum_by_year[yr] = cum_by_year.get(yr, 0) + row["Prévision (Mds)"]
        new_totals.append(round(cum_by_year[yr], 2))
    full["Total cumulé"] = new_totals

    return full, by_year


# =============================================================================
# 4. KPIs HISTORIQUE & PRÉVISION
# =============================================================================
def compute_kpis(series: pd.Series) -> dict:
    """KPIs pour l'onglet visualisation historique."""
    s = series.dropna()
    if len(s) == 0:
        return {"last_date": "-", "last_value": 0, "max_date": "-", "max_value": 0,
                "var_mom": 0, "var_yoy": 0,
                "mean_12m": 0, "cv_12m": 0, "last_year": 0, "total_last_year": 0,
                "yoy_total": 0}

    if hasattr(s.index, "to_timestamp"):
        s.index = s.index.to_timestamp()

    last_value = float(s.iloc[-1])
    last_date  = s.index[-1].strftime("%b %Y") if hasattr(s.index[-1], "strftime") else str(s.index[-1])
    last_year  = int(s.index[-1].year) if hasattr(s.index[-1], "year") else 0

    # Valeur maximale et sa date
    max_idx = s.idxmax()
    max_value = float(s.loc[max_idx])
    max_date  = max_idx.strftime("%b %Y") if hasattr(max_idx, "strftime") else str(max_idx)

    var_mom = float((s.iloc[-1] / s.iloc[-2] - 1) * 100) if len(s) >= 2 else 0.0
    var_yoy = float((s.iloc[-1] / s.iloc[-13] - 1) * 100) if len(s) >= 13 else 0.0

    last12 = s.iloc[-12:] if len(s) >= 12 else s
    mean_12m = float(last12.mean())
    cv_12m   = float((last12.std() / last12.mean()) * 100) if last12.mean() else 0.0

    # Total annuel dernière année & variation totale vs N-1
    total_last_year = float(s[s.index.year == last_year].sum())
    prev_year_data  = s[s.index.year == last_year - 1]
    total_prev = float(prev_year_data.sum()) if len(prev_year_data) > 0 else 0.0
    yoy_total = ((total_last_year / total_prev - 1) * 100) if total_prev else 0.0

    return {
        "last_date":       last_date,
        "last_value":      last_value,
        "max_date":        max_date,
        "max_value":       max_value,
        "var_mom":         var_mom,
        "var_yoy":         var_yoy,
        "mean_12m":        mean_12m,
        "cv_12m":          cv_12m,
        "last_year":       last_year,
        "total_last_year": total_last_year,
        "yoy_total":       yoy_total,
    }


def compute_forecast_kpis(prev_df: pd.DataFrame, df_hist: pd.DataFrame,
                          horizon: int, metrics_df=None) -> dict:
    """KPIs spécifiques à l'onglet prévision."""
    if prev_df is None or prev_df.empty:
        return {"total": 0, "growth": 0, "mape": 0, "ic_width": 0, "ref_year": 0}

    serie = df_hist["Recettes_fiscales"].dropna()
    if hasattr(serie.index, "to_timestamp"):
        serie.index = serie.index.to_timestamp()

    total = float(prev_df["Prévision (Mds)"].sum())

    # Croissance attendue : on compare au cumul sur la même longueur à la fin de l'historique
    n_periods = len(prev_df)
    last_periods = serie.iloc[-n_periods:].sum() if len(serie) >= n_periods else serie.sum()
    growth = ((total / last_periods - 1) * 100) if last_periods else 0.0
    ref_year = int(serie.index[-1].year) if len(serie) else 0

    # MAPE du meilleur modèle sur le test
    mape = 0.0
    if metrics_df is not None:
        try:
            best_model = "RandomForest" if horizon == 12 else "XGBoost"
            mape = float(metrics_df[metrics_df["Modele"] == best_model]["MAPE"].iloc[0])
        except Exception:
            try:
                mape = float(metrics_df["MAPE"].min())
            except Exception:
                mape = 0.0

    # Largeur moyenne des IC (demi-largeur)
    ic_width = float(((prev_df["IC haut 95%"] - prev_df["IC bas 95%"]) / 2).mean())

    return {
        "total":    total,
        "growth":   growth,
        "mape":     mape,
        "ic_width": ic_width,
        "ref_year": ref_year,
    }


# =============================================================================
# 5. GRAPHIQUES PNG POUR LES RAPPORTS
# =============================================================================
def _setup_mpl_style():
    plt.rcParams.update({
        "font.family":     "DejaVu Sans",
        "axes.labelweight":"bold",
        "axes.titleweight":"bold",
        "axes.spines.top": False,
        "axes.spines.right":False,
        "axes.edgecolor":  "#1B2559",
        "axes.labelcolor": "#1B2559",
        "xtick.color":     "#1B2559",
        "ytick.color":     "#1B2559",
        "font.size":       10,
        "axes.labelsize":  11,
        "axes.titlesize":  12,
        "xtick.labelsize": 10,
        "ytick.labelsize": 10,
    })


def _chart_history(df_hist: pd.DataFrame) -> bytes:
    _setup_mpl_style()
    fig, ax = plt.subplots(figsize=(9, 4.6), dpi=130)
    s = df_hist["Recettes_fiscales"].dropna()
    ax.fill_between(s.index, 0, s.values, color=PURPLE, alpha=0.10)
    ax.plot(s.index, s.values, color=PURPLE, linewidth=2.0, label="Recettes fiscales")
    ma = s.rolling(12).mean()
    ax.plot(ma.index, ma.values, color=ACCENT_GOLD, linewidth=2.0, linestyle="--", label="Moyenne mobile 12 mois")
    ax.set_title("Évolution mensuelle des recettes fiscales (2010 à 2024)")
    ax.set_xlabel("Date"); ax.set_ylabel("Mds FCFA")
    ax.grid(True, alpha=0.25, color="#A3AED0")
    ax.legend(loc="upper left", frameon=False)
    buf = io.BytesIO(); fig.tight_layout(); fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _chart_forecast(prev_df: pd.DataFrame, horizon: int) -> bytes:
    _setup_mpl_style()
    fig, ax = plt.subplots(figsize=(9, 4.6), dpi=130)
    dates = pd.to_datetime(prev_df["Mois"] + "-01")
    ax.plot(dates, prev_df["Prévision (Mds)"], color=PURPLE, marker="D", markersize=6,
            linewidth=2.2, label=f"Prévision (H = {horizon})")
    ax.fill_between(dates, prev_df["IC bas 95%"], prev_df["IC haut 95%"],
                    color=PURPLE, alpha=0.18, label="Intervalle de confiance 95 %")
    model_name = "RandomForest" if horizon == 12 else "XGBoost"
    ax.set_title(f"Prévisions des recettes fiscales - Horizon {horizon} mois ({model_name})")
    ax.set_xlabel("Date"); ax.set_ylabel("Mds FCFA")
    ax.grid(True, alpha=0.25, color="#A3AED0")
    ax.legend(loc="upper left", frameon=False)
    plt.xticks(rotation=30, ha="right")
    buf = io.BytesIO(); fig.tight_layout(); fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _chart_validation(bt: pd.DataFrame, horizon: int) -> bytes:
    _setup_mpl_style()
    fig, ax = plt.subplots(figsize=(9, 4.3), dpi=130)
    x = np.arange(len(bt))
    w = 0.4
    ax.bar(x - w/2, bt["Réel (Mds)"],  width=w, color=NAVY,      label="Réel 2024")
    ax.bar(x + w/2, bt["Prévu (Mds)"], width=w, color=PURPLE_LT, label="Prévu par le modèle")
    ax.set_xticks(x); ax.set_xticklabels(bt["Mois"], rotation=30, ha="right")
    ax.set_title(f"Validation : Prévu vs Réel sur 2024 - Horizon {horizon}")
    ax.set_xlabel("Mois"); ax.set_ylabel("Mds FCFA")
    ax.grid(True, alpha=0.25, color="#A3AED0", axis="y")
    ax.legend(frameon=False)
    buf = io.BytesIO(); fig.tight_layout(); fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# =============================================================================
# 6. RAPPORT WORD
# =============================================================================
def generate_word_report(*, titre, horizon, n_years, kpis, metrics_d, prev_df,
                         by_year, backtest, fkpis, include_kpis, include_hist,
                         include_prev, include_val, include_mtr, include_shap,
                         include_meth, df_hist, logo_path, figures_dir,
                         **_) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Réglage marges
    for section in doc.sections:
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(2.0)

    # === ENTÊTE INSTITUTIONNEL =============================================
    header_para = doc.sections[0].header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tableau d'en-tête (logo + texte centré)
    head_tbl = doc.add_table(rows=1, cols=3)
    head_tbl.autofit = False
    head_tbl.columns[0].width = Cm(3.0)
    head_tbl.columns[1].width = Cm(11.0)
    head_tbl.columns[2].width = Cm(3.0)

    # Cellule logo
    if Path(logo_path).exists():
        cell_logo = head_tbl.cell(0, 0)
        cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Cm(2.5))

    # Cellule texte central
    cell_mid = head_tbl.cell(0, 1)
    cell_mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for txt, size, bold, color in [
        ("RÉPUBLIQUE DU CAMEROUN", 10, True,  RGBColor(0x0F, 0x15, 0x35)),
        ("Paix · Travail · Patrie", 9,  False, RGBColor(0x00, 0x00, 0x00)),
        ("",                        8,  False, None),
        ("MINISTÈRE DES FINANCES",  12, True,  RGBColor(0x43, 0x18, 0xFF)),
        ("Ministry of Finance",     9,  False, RGBColor(0x00, 0x00, 0x00)),
    ]:
        p = cell_mid.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if txt:
            r = p.add_run(txt)
            r.font.size = Pt(size)
            r.font.bold = bold
            if color: r.font.color.rgb = color

    # === SÉPARATEUR ========================================================
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("─" * 90)
    r_sep.font.color.rgb = RGBColor(0x43, 0x18, 0xFF)
    r_sep.font.size = Pt(8)

    # === TITRE PRINCIPAL ===================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(titre)
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sub_txt = f"Horizon de prévision : {horizon} mois"
    if horizon == 12 and n_years > 1:
        sub_txt += f" projeté sur {n_years} ans"
    best_name = "Random Forest" if horizon == 12 else "XGBoost"
    sub_txt += f"  ·  Modèle {best_name}"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(sub_txt)
    r_sub.font.size = Pt(12); r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run(f"Édité le {datetime.now().strftime('%d / %m / %Y')}")
    r_date.font.size = Pt(10)
    r_date.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # === 1. SYNTHÈSE EXÉCUTIVE ============================================
    doc.add_heading("1. Synthèse exécutive", level=1)
    p = doc.add_paragraph(
        "Ce rapport présente les résultats de l'analyse et de la prévision des "
        "recettes fiscales du Cameroun, produits par le Système d'Analyse et de "
        "Prévision développé sur la base de techniques avancées de Machine Learning. "
        "Six familles de modèles (Ridge, Lasso, ElasticNet, Random Forest, XGBoost et SVR) "
        "ont été évaluées. XGBoost a été retenu pour les horizons trimestriel et semestriel, "
        "et Random Forest pour l'horizon annuel, chacun dominant sur les principales "
        "métriques de performance (RMSE, MAE, MAPE, R²) pour son horizon respectif."
    )

    # === 2. INDICATEURS CLÉS ==============================================
    sec_num = 2
    if include_kpis and kpis:
        doc.add_heading(f"{sec_num}. Indicateurs clés observés", level=1)
        sec_num += 1

        rows = [
            ("Dernière valeur mensuelle observée", f"{kpis['last_value']:.2f} Mds FCFA"),
            ("Période concernée",                  f"{kpis['last_date']}"),
            ("Total annuel " + str(kpis['last_year']), f"{kpis['total_last_year']:,.0f} Mds FCFA"),
            ("Variation annuelle vs N-1",          f"{kpis['yoy_total']:+.2f} %"),
            ("Moyenne mensuelle sur 12 mois",      f"{kpis['mean_12m']:.2f} Mds FCFA"),
            ("Coefficient de variation 12 mois",   f"{kpis['cv_12m']:.1f} %"),
        ]
        _make_kv_table(doc, rows, header_color=NAVY)

    # === 3. ÉVOLUTION HISTORIQUE ==========================================
    if include_hist and df_hist is not None and "Recettes_fiscales" in df_hist.columns:
        doc.add_heading(f"{sec_num}. Évolution historique", level=1)
        sec_num += 1
        img = _chart_history(df_hist)
        doc.add_picture(io.BytesIO(img), width=Inches(6.3))
        _caption(doc, "Figure 1 - Évolution mensuelle des recettes fiscales (2010 à 2024)")

    # === 4. COMPARAISON MODÈLES ===========================================
    if include_mtr and metrics_d is not None:
        doc.add_heading(f"{sec_num}. Performance comparée des modèles (H = {horizon})", level=1)
        sec_num += 1
        _make_df_table(doc, metrics_d.round(3), header_color=NAVY)
        doc.add_paragraph(
            "Le modèle retenu pour cet horizon domine sur les principales métriques "
            "de performance ce qui justifie sa sélection."
        )

    # === 5. PRÉVISIONS ====================================================
    if include_prev and prev_df is not None and fkpis is not None:
        doc.add_heading(f"{sec_num}. Prévisions sur {horizon} mois", level=1)
        sec_num += 1

        rows = [
            ("Total prévu sur la période",       f"{fkpis['total']:,.0f} Mds FCFA"),
            ("Croissance attendue",              f"{fkpis['growth']:+.2f} % vs même période {fkpis['ref_year']}"),
            ("Précision attendue du modèle",     f"MAPE = {fkpis['mape']:.2f} %"),
        ]
        _make_kv_table(doc, rows, header_color=PURPLE)

        if by_year is not None and len(by_year) > 1:
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run("Cumul prévu par année")
            r.font.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            for yr, total in by_year.items():
                doc.add_paragraph(f"  •  Année {yr} : {total:,.0f} Mds FCFA", style="List Bullet")

        doc.add_paragraph()
        img = _chart_forecast(prev_df, horizon)
        doc.add_picture(io.BytesIO(img), width=Inches(6.3))
        _caption(doc, f"Figure - Prévisions mensuelles avec intervalle de confiance 95 %")

        doc.add_paragraph()
        _make_df_table(doc, prev_df.round(2), header_color=PURPLE)

    # === 6. VALIDATION 2024 ===============================================
    if include_val and backtest is not None:
        doc.add_heading(f"{sec_num}. Validation sur l'année 2024", level=1)
        sec_num += 1
        mape_bt = backtest["Écart (%)"].abs().mean()
        doc.add_paragraph(
            f"Confrontation des prévisions aux valeurs réellement observées en 2024. "
            f"L'erreur moyenne absolue en pourcentage (MAPE) constatée s'élève à {mape_bt:.2f} %."
        )
        img = _chart_validation(backtest, horizon)
        doc.add_picture(io.BytesIO(img), width=Inches(6.3))
        _caption(doc, "Figure - Comparaison Réel vs Prévu sur 2024")
        _make_df_table(doc, backtest.round(2), header_color=ACCENT_GREEN)

    # === 7. SHAP ===========================================================
    if include_shap:
        doc.add_heading(f"{sec_num}. Variables les plus contributives (SHAP)", level=1)
        sec_num += 1
        doc.add_paragraph(
            "L'analyse SHAP (SHapley Additive exPlanations) quantifie la contribution de "
            "chaque variable explicative à la prévision produite par le modèle XGBoost. "
            "Plus la valeur SHAP est élevée, plus la variable influence la sortie du modèle."
        )
        shap_path = Path(figures_dir) / f"11_SHAP_XGBoost_H{horizon}.png"
        if shap_path.exists():
            doc.add_picture(str(shap_path), width=Inches(6.3))
            _caption(doc, f"Figure - Valeurs SHAP pour le modèle XGBoost H = {horizon}")

    # === 8. MÉTHODOLOGIE ==================================================
    if include_meth:
        doc.add_heading(f"{sec_num}. Méthodologie", level=1)
        sec_num += 1

        for titre_p, contenu in [
            ("Données utilisées",
             "Recettes fiscales mensuelles du Cameroun de janvier 2010 à décembre 2024 "
             "(180 observations) accompagnées de 14 variables exogènes macroéconomiques "
             "et climatiques (PIB, IPI, Inflation, Prix Brent, REER, Dépenses publiques, "
             "Solde budgétaire, prix internationaux des matières premières, précipitations, etc.)."),
            ("Ingénierie des features",
             "Construction de lags de la série cible, décomposition STL (tendance + saisonnalité + résidu), "
             "sélection des lags résiduels les plus informatifs par corrélation croisée, agrégation "
             "trimestrielle pour les variables disponibles uniquement à fréquence basse."),
            ("Famille de modèles évalués",
             "Six modèles : Ridge, Lasso, ElasticNet (linéaires régularisés), Random Forest, "
             "XGBoost (ensemble) et SVR (noyau RBF). Validation croisée temporelle pour éviter "
             "toute fuite d'information future."),
            ("Critère de sélection",
             "XGBoost a été retenu pour les horizons H=3 et H=6, et Random Forest pour H=12. Chacun "
             "domine sur les principales métriques de performance (RMSE, MAE, MAPE, R²) pour son horizon respectif."),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(titre_p); r.font.bold = True
            doc.add_paragraph(contenu)

    # === PIED DE PAGE ======================================================
    footer = doc.sections[0].footer
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(
        "Document généré par le Système d'Analyse et de Prévision des Recettes Fiscales · MINFI Cameroun"
    )
    rf.font.size = Pt(8); rf.font.italic = True
    rf.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def _caption(doc, text):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _make_kv_table(doc, rows, header_color="#4318FF"):
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True
    for i, (k, v) in enumerate(rows):
        c1 = tbl.cell(i, 0); c2 = tbl.cell(i, 1)
        c1.text = ""; c2.text = ""
        r1 = c1.paragraphs[0].add_run(k); r1.font.bold = True; r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(v); r2.font.size = Pt(10)


def _make_df_table(doc, df: pd.DataFrame, header_color="#4318FF"):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    cols = df.columns.tolist()
    tbl = doc.add_table(rows=len(df) + 1, cols=len(cols))
    tbl.style = "Light Grid Accent 1"

    # En-tête
    for j, c in enumerate(cols):
        cell = tbl.cell(0, j); cell.text = ""
        r = cell.paragraphs[0].add_run(str(c)); r.font.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # background
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), header_color.replace("#", ""))
        tcPr.append(shd)

    # Données
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, v in enumerate(row):
            cell = tbl.cell(i, j); cell.text = ""
            r = cell.paragraphs[0].add_run(str(v)); r.font.size = Pt(9)


# =============================================================================
# 7. RAPPORT PDF  — style LaTeX professionnel
# =============================================================================
def generate_pdf_report(*, titre, horizon, n_years, kpis, metrics_d, prev_df,
                        by_year, backtest, fkpis, include_kpis, include_hist,
                        include_prev, include_val, include_mtr, include_shap,
                        include_meth, df_hist, logo_path, figures_dir,
                        **_) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
        PageBreak, KeepTogether, HRFlowable,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT

    buf = io.BytesIO()

    # ── palette ──────────────────────────────────────────────────────────
    NAVY_C    = colors.HexColor(NAVY)
    PURPLE_C  = colors.HexColor(PURPLE)
    BLACK_C   = colors.black
    GRAY_MID  = colors.HexColor("#888888")
    RULE_GRAY = colors.HexColor("#AAAAAA")
    ALT_ROW   = colors.HexColor("#F4F4F4")   # ligne alternée légère

    PAGE_W, PAGE_H = A4
    L_MAR  = 2.5 * cm
    R_MAR  = 2.5 * cm
    T_MAR  = 2.0 * cm
    B_MAR  = 2.0 * cm
    CW     = PAGE_W - L_MAR - R_MAR        # largeur utile ≈ 16.5 cm

    # ── en-tête / pied de page (toutes pages) ────────────────────────────
    def _header_footer(canvas, doc_):
        canvas.saveState()
        # Bande navy
        canvas.setFillColor(NAVY_C)
        canvas.rect(0, PAGE_H - 1.1 * cm, PAGE_W, 1.1 * cm, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont("Times-Bold", 9)
        canvas.drawString(L_MAR, PAGE_H - 0.73 * cm,
                          "MINISTÈRE DES FINANCES · CAMEROUN")
        canvas.setFont("Times-Roman", 8)
        canvas.drawRightString(PAGE_W - R_MAR, PAGE_H - 0.73 * cm,
                               "Système d'Analyse et de Prévision des Recettes Fiscales")
        # Filet bas
        canvas.setStrokeColor(GRAY_MID)
        canvas.setLineWidth(0.4)
        canvas.line(L_MAR, 1.35 * cm, PAGE_W - R_MAR, 1.35 * cm)
        # Numéro de page
        canvas.setFillColor(BLACK_C)
        canvas.setFont("Times-Italic", 9)
        canvas.drawCentredString(PAGE_W / 2, 0.8 * cm, f"Page {doc_.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=L_MAR, rightMargin=R_MAR,
        topMargin=T_MAR, bottomMargin=B_MAR,
    )

    # ══════════════════════════════════════════════════════════════════════
    # STYLES TYPOGRAPHIQUES  — Times New Roman 12 pt, texte justifié
    # ══════════════════════════════════════════════════════════════════════
    BODY_SIZE = 12
    LEAD      = 18          # interligne LaTeX ≈ 1.5 × corps

    cover_repub = ParagraphStyle(
        "CRep", fontName="Times-Bold", fontSize=13, textColor=NAVY_C,
        alignment=TA_CENTER, leading=18, spaceAfter=3,
    )
    cover_devise = ParagraphStyle(
        "CDev", fontName="Times-Italic", fontSize=11, textColor=BLACK_C,
        alignment=TA_CENTER, leading=15, spaceAfter=10,
    )
    cover_minfi = ParagraphStyle(
        "CMF", fontName="Times-Bold", fontSize=18, textColor=PURPLE_C,
        alignment=TA_CENTER, leading=24, spaceAfter=3,
    )
    cover_eng = ParagraphStyle(
        "CEng", fontName="Times-Italic", fontSize=11, textColor=BLACK_C,
        alignment=TA_CENTER, leading=15, spaceAfter=6,
    )
    cover_title = ParagraphStyle(
        "CTit", fontName="Times-Bold", fontSize=24, textColor=BLACK_C,
        alignment=TA_CENTER, leading=32, spaceAfter=14,
    )
    cover_sub = ParagraphStyle(
        "CSub", fontName="Times-Italic", fontSize=12, textColor=BLACK_C,
        alignment=TA_CENTER, leading=17, spaceAfter=5,
    )
    # Titres de section : gras 14 pt, filet sous-jacent géré séparément
    h1_st = ParagraphStyle(
        "H1", fontName="Times-Bold", fontSize=14, textColor=BLACK_C,
        leading=20, spaceBefore=22, spaceAfter=6, leftIndent=0,
    )
    # Sous-section (méthodologie)
    h2_st = ParagraphStyle(
        "H2", fontName="Times-BoldItalic", fontSize=12, textColor=BLACK_C,
        leading=17, spaceBefore=12, spaceAfter=4, leftIndent=0,
    )
    # Corps principal : Times 12 pt, justifié, interligne 18 pt
    body_st = ParagraphStyle(
        "BD", fontName="Times-Roman", fontSize=BODY_SIZE, textColor=BLACK_C,
        alignment=TA_JUSTIFY, leading=LEAD, spaceBefore=0, spaceAfter=8,
    )
    # Corps gras (liste de puces)
    bold_body = ParagraphStyle(
        "BB", fontName="Times-Bold", fontSize=BODY_SIZE, textColor=BLACK_C,
        leading=LEAD, spaceBefore=6, spaceAfter=4,
    )
    # Légende de figure : italic 10 pt, centré
    caption_st = ParagraphStyle(
        "CAP", fontName="Times-Italic", fontSize=10, textColor=GRAY_MID,
        alignment=TA_CENTER, leading=14, spaceBefore=5, spaceAfter=14,
    )
    # Note de bas de tableau
    note_st = ParagraphStyle(
        "NOTE", fontName="Times-Italic", fontSize=10, textColor=GRAY_MID,
        leading=14, spaceBefore=4, spaceAfter=10,
    )

    # ══════════════════════════════════════════════════════════════════════
    # HELPERS  TABLEAUX — style booktabs (pas de bordures verticales)
    # ══════════════════════════════════════════════════════════════════════
    TOPRULE    = 1.5
    MIDRULE    = 0.75
    BOTTOMRULE = 1.5
    CMIDRULE   = 0.3          # filet léger entre lignes de données

    def _booktabs_style(n_data_rows, bold_row=None, alt=True,
                        font_size=11, padding=7, align_right_cols=None):
        """
        Renvoie un TableStyle booktabs :
          - toprule (1.5pt), midrule (0.75pt), bottomrule (1.5pt)
          - pas de lignes verticales
          - filet léger entre lignes de données
          - alternance optionnelle de fond
        """
        cmds = [
            ("LINEABOVE",  (0, 0),  (-1, 0),  TOPRULE,    BLACK_C),
            ("LINEBELOW",  (0, 0),  (-1, 0),  MIDRULE,    BLACK_C),
            ("LINEBELOW",  (0, -1), (-1, -1), BOTTOMRULE, BLACK_C),
            ("FONTNAME",   (0, 0),  (-1, 0),  "Times-Bold"),
            ("FONTSIZE",   (0, 0),  (-1, -1), font_size),
            ("TEXTCOLOR",  (0, 0),  (-1, -1), BLACK_C),
            ("TOPPADDING",    (0, 0), (-1, -1), padding),
            ("BOTTOMPADDING", (0, 0), (-1, -1), padding),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
            ("VALIGN",     (0, 0),  (-1, -1), "MIDDLE"),
        ]
        # Filets entre lignes de données
        for r in range(1, n_data_rows):
            cmds.append(("LINEBELOW", (0, r), (-1, r), CMIDRULE, RULE_GRAY))
        # Alternance fond
        if alt:
            for r in range(1, n_data_rows + 1, 2):
                cmds.append(("BACKGROUND", (0, r), (-1, r), ALT_ROW))
        # Colonnes numériques alignées à droite
        if align_right_cols:
            for c in align_right_cols:
                cmds.append(("ALIGN", (c, 0), (c, -1), "RIGHT"))
        # Ligne du meilleur modèle en gras
        if bold_row is not None:
            cmds.append(("FONTNAME", (0, bold_row), (-1, bold_row), "Times-Bold"))
        return TableStyle(cmds)

    def _make_table(data, col_widths, bold_row=None, alt=True,
                    font_size=11, padding=7, align_right_cols=None):
        n_data = len(data) - 1
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(_booktabs_style(n_data, bold_row=bold_row, alt=alt,
                                   font_size=font_size, padding=padding,
                                   align_right_cols=align_right_cols))
        return t

    def _kv_table(rows_data, label="Indicateur", value_label="Valeur",
                  left_frac=0.60):
        """Tableau clé / valeur (2 colonnes) — valeur en gras à droite."""
        lw = CW * left_frac
        rw = CW - lw
        data = [[label, value_label]] + [[k, v] for k, v in rows_data]
        n_data = len(data) - 1
        cmds = [
            ("LINEABOVE",  (0, 0),  (-1, 0),  TOPRULE,    BLACK_C),
            ("LINEBELOW",  (0, 0),  (-1, 0),  MIDRULE,    BLACK_C),
            ("LINEBELOW",  (0, -1), (-1, -1), BOTTOMRULE, BLACK_C),
            ("FONTNAME",   (0, 0),  (-1, 0),  "Times-Bold"),
            ("FONTNAME",   (1, 1),  (1,  -1), "Times-Bold"),
            ("FONTSIZE",   (0, 0),  (-1, -1), 11),
            ("TEXTCOLOR",  (0, 0),  (-1, -1), BLACK_C),
            ("ALIGN",      (1, 0),  (1,  -1), "RIGHT"),
            ("TOPPADDING",    (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
            ("VALIGN",     (0, 0),  (-1, -1), "MIDDLE"),
        ]
        for r in range(1, n_data):
            cmds.append(("LINEBELOW", (0, r), (-1, r), CMIDRULE, RULE_GRAY))
        for r in range(1, n_data + 1, 2):
            cmds.append(("BACKGROUND", (0, r), (-1, r), ALT_ROW))
        t = Table(data, colWidths=[lw, rw])
        t.setStyle(TableStyle(cmds))
        return t

    def _img_c(src, w=None, h=None):
        """Image centrée."""
        img = Image(src, width=w or CW, height=h or 8 * cm)
        img.hAlign = "CENTER"
        return img

    def _section_rule():
        """Filet fin sous un titre de section (style LaTeX article)."""
        return HRFlowable(width="100%", thickness=0.5, color=RULE_GRAY,
                          spaceAfter=6, spaceBefore=0)

    story = []

    # ══════════════════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ══════════════════════════════════════════════════════════════════════
    story.append(Spacer(1, 1.0 * cm))

    # République du Cameroun + devise
    story.append(Paragraph("RÉPUBLIQUE DU CAMEROUN", cover_repub))
    story.append(Paragraph("<i>Paix · Travail · Patrie</i>", cover_devise))

    # Logo MINFI centré juste au-dessus de "Ministère des Finances"
    if Path(logo_path).exists():
        logo = Image(logo_path, width=5.0 * cm, height=5.0 * cm)
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 0.4 * cm))

    # Ministère des Finances
    story.append(Paragraph("MINISTÈRE DES FINANCES", cover_minfi))
    story.append(Paragraph("<i>Ministry of Finance</i>", cover_eng))

    # Filet décoratif double
    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=2.0, color=PURPLE_C,
                             spaceAfter=2, spaceBefore=0))
    story.append(HRFlowable(width="100%", thickness=0.5, color=PURPLE_C,
                             spaceAfter=0, spaceBefore=4))
    story.append(Spacer(1, 2.0 * cm))

    # Titre du rapport
    story.append(Paragraph(titre, cover_title))

    best_name = "Random Forest" if horizon == 12 else "XGBoost"
    sub_txt   = f"Horizon de prévision : {horizon} mois"
    if horizon == 12 and n_years > 1:
        sub_txt += f" projeté sur {n_years} ans"
    sub_txt += f" — Modèle {best_name}"
    story.append(Paragraph(sub_txt, cover_sub))
    story.append(Paragraph(
        f"Édité le {datetime.now().strftime('%d / %m / %Y')}",
        cover_sub,
    ))

    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════════════
    # 1.  SYNTHÈSE EXÉCUTIVE
    # ══════════════════════════════════════════════════════════════════════
    story.append(Paragraph("1. Synthèse exécutive", h1_st))
    story.append(_section_rule())
    story.append(Paragraph(
        "Ce rapport présente les résultats de l'analyse et de la prévision des recettes "
        "fiscales du Cameroun, produits par le Système d'Analyse et de Prévision développé "
        "sur la base de techniques avancées de Machine Learning. Six familles de modèles "
        "(Ridge, Lasso, ElasticNet, Random Forest, XGBoost et SVR) ont été évaluées. "
        "XGBoost a été retenu pour les horizons trimestriel et semestriel, et Random Forest "
        "pour l'horizon annuel, chacun dominant sur les principales métriques de performance "
        "(RMSE, MAE, MAPE, R²) pour son horizon respectif.",
        body_st,
    ))

    sec_num = 2

    # ══════════════════════════════════════════════════════════════════════
    # 2.  INDICATEURS CLÉS
    # ══════════════════════════════════════════════════════════════════════
    if include_kpis and kpis:
        story.append(Paragraph(f"{sec_num}. Indicateurs clés observés", h1_st))
        story.append(_section_rule())
        sec_num += 1
        kpi_rows = [
            ("Dernière valeur mensuelle observée", f"{kpis['last_value']:.2f} Mds FCFA"),
            ("Période concernée",                  str(kpis["last_date"])),
            (f"Total annuel {kpis['last_year']}",  f"{kpis['total_last_year']:,.0f} Mds FCFA"),
            ("Variation annuelle vs N‑1",      f"{kpis['yoy_total']:+.2f} %"),
            ("Moyenne mensuelle sur 12 mois",       f"{kpis['mean_12m']:.2f} Mds FCFA"),
            ("Coefficient de variation 12 mois",    f"{kpis['cv_12m']:.1f} %"),
        ]
        story.append(_kv_table(kpi_rows))
        story.append(Spacer(1, 8))

    # ══════════════════════════════════════════════════════════════════════
    # 3.  ÉVOLUTION HISTORIQUE
    # ══════════════════════════════════════════════════════════════════════
    if include_hist and df_hist is not None and "Recettes_fiscales" in df_hist.columns:
        story.append(Paragraph(f"{sec_num}. Évolution historique", h1_st))
        story.append(_section_rule())
        sec_num += 1
        img_bytes = _chart_history(df_hist)
        story.append(_img_c(io.BytesIO(img_bytes), CW, 8.5 * cm))
        story.append(Paragraph(
            "Figure 1 — Évolution mensuelle des recettes fiscales "
            "(2010 à 2024).",
            caption_st,
        ))

    # ══════════════════════════════════════════════════════════════════════
    # 4.  PERFORMANCE DES MODÈLES
    # ══════════════════════════════════════════════════════════════════════
    if include_mtr and metrics_d is not None:
        story.append(Paragraph(
            f"{sec_num}. Performance comparée des modèles "
            f"(H = {horizon} mois)", h1_st,
        ))
        story.append(_section_rule())
        sec_num += 1

        df_m  = metrics_d.copy()
        # arrondi soigné colonne par colonne
        for col in df_m.columns:
            if col != "Modele":
                try:
                    df_m[col] = df_m[col].astype(float).round(3)
                except Exception:
                    pass

        cols = df_m.columns.tolist()
        rows = [cols] + df_m.astype(str).values.tolist()

        best_model_name = "XGBoost" if horizon in (3, 6) else "RandomForest"
        try:
            best_row_idx = int(df_m[df_m["Modele"] == best_model_name].index[0]) + 1
        except Exception:
            try:
                best_row_idx = int(df_m["RMSE"].astype(float).idxmin()) + 1
            except Exception:
                best_row_idx = None

        n_c    = len(cols)
        cw0    = 4.5 * cm
        cw_r   = (CW - cw0) / max(n_c - 1, 1)
        cws    = [cw0] + [cw_r] * (n_c - 1)
        rc     = list(range(1, n_c))    # colonnes numériques

        story.append(_make_table(rows, cws, bold_row=best_row_idx,
                                 font_size=11, padding=7, align_right_cols=rc))
        story.append(Paragraph(
            "<i>Note : la ligne en gras indique le modèle retenu pour cet horizon "
            "de prévision sur la base de la convergence des métriques RMSE, MAE, MAPE et R².</i>",
            note_st,
        ))

    # ══════════════════════════════════════════════════════════════════════
    # 5.  PRÉVISIONS
    # ══════════════════════════════════════════════════════════════════════
    if include_prev and prev_df is not None and fkpis is not None:
        story.append(Paragraph(
            f"{sec_num}. Prévisions sur {horizon} mois", h1_st,
        ))
        story.append(_section_rule())
        sec_num += 1

        # Tableau synthèse de prévision
        prev_rows = [
            ("Total prévu sur la période",
             f"{fkpis['total']:,.0f} Mds FCFA"),
            ("Croissance attendue",
             f"{fkpis['growth']:+.2f} % vs même période {fkpis['ref_year']}"),
            ("Précision attendue du modèle (MAPE test)",
             f"{fkpis['mape']:.2f} %"),
        ]
        story.append(_kv_table(prev_rows,
                               label="Indicateur de prévision", value_label="Valeur"))
        story.append(Spacer(1, 10))

        # Cumul par année
        if by_year is not None and len(by_year) > 1:
            story.append(Paragraph("<b>Cumul prévu par année</b>", bold_body))
            for yr, total in by_year.items():
                story.append(Paragraph(
                    f"  •  Année <b>{yr}</b> : "
                    f"<b>{total:,.0f} Mds FCFA</b>",
                    body_st,
                ))
            story.append(Spacer(1, 10))

        # Graphique
        img_bytes = _chart_forecast(prev_df, horizon)
        story.append(_img_c(io.BytesIO(img_bytes), CW, 8.5 * cm))
        story.append(Paragraph(
            "Figure 2 — Prévisions mensuelles des recettes fiscales "
            "avec intervalle de confiance à 95 %.",
            caption_st,
        ))

        # Tableau détail
        df_p   = prev_df.round(2)
        data_p = [df_p.columns.tolist()] + df_p.astype(str).values.tolist()
        n_c    = len(df_p.columns)
        cws    = [CW / n_c] * n_c
        rc     = list(range(1, n_c))
        story.append(_make_table(data_p, cws, font_size=10, padding=5,
                                 align_right_cols=rc))

    # ══════════════════════════════════════════════════════════════════════
    # 6.  VALIDATION 2024
    # ══════════════════════════════════════════════════════════════════════
    if include_val and backtest is not None:
        story.append(Paragraph(
            f"{sec_num}. Validation sur l'année 2024", h1_st,
        ))
        story.append(_section_rule())
        sec_num += 1

        mape_bt = backtest["Écart (%)"].abs().mean()
        story.append(Paragraph(
            "La confrontation des prévisions du modèle aux valeurs réellement observées "
            "sur l'année 2024 permet d'évaluer la qualité opérationnelle des prévisions. "
            f"L'erreur moyenne absolue en pourcentage (MAPE) constatée sur cette période "
            f"s'élève à <b>{mape_bt:.2f} %</b>.",
            body_st,
        ))
        story.append(Spacer(1, 6))

        img_bytes = _chart_validation(backtest, horizon)
        story.append(_img_c(io.BytesIO(img_bytes), CW, 8.0 * cm))
        story.append(Paragraph(
            "Figure 3 — Comparaison Réel vs Prévu sur 2024 "
            f"(horizon H = {horizon} mois).",
            caption_st,
        ))

        df_b   = backtest.round(2)
        data_b = [df_b.columns.tolist()] + df_b.astype(str).values.tolist()
        n_c    = len(df_b.columns)
        cws    = [CW / n_c] * n_c
        rc     = list(range(1, n_c))
        story.append(_make_table(data_b, cws, font_size=10, padding=5,
                                 align_right_cols=rc))

    # ══════════════════════════════════════════════════════════════════════
    # 7.  SHAP
    # ══════════════════════════════════════════════════════════════════════
    if include_shap:
        story.append(Paragraph(
            f"{sec_num}. Variables les plus contributives (SHAP)", h1_st,
        ))
        story.append(_section_rule())
        sec_num += 1

        story.append(Paragraph(
            "L'analyse SHAP (SHapley Additive exPlanations) quantifie la contribution "
            "marginale de chaque variable explicative à la prévision produite par le modèle "
            "XGBoost. La valeur SHAP absolue moyenne mesure l'influence globale d'une variable "
            "sur l'ensemble du jeu de test : plus cette valeur est élevée, plus la "
            "variable oriente significativement la sortie du modèle.",
            body_st,
        ))
        shap_path = Path(figures_dir) / f"11_SHAP_XGBoost_H{horizon}.png"
        if shap_path.exists():
            story.append(Spacer(1, 6))
            story.append(_img_c(str(shap_path), CW, 11 * cm))
            story.append(Paragraph(
                f"Figure 4 — Importance des variables selon les valeurs "
                f"SHAP — modèle XGBoost (H = {horizon} mois).",
                caption_st,
            ))

    # ══════════════════════════════════════════════════════════════════════
    # 8.  MÉTHODOLOGIE
    # ══════════════════════════════════════════════════════════════════════
    if include_meth:
        story.append(Paragraph(f"{sec_num}. Méthodologie", h1_st))
        story.append(_section_rule())

        meth_sections = [
            ("Données utilisées",
             "Les données mobilisées comprennent les recettes fiscales mensuelles du "
             "Cameroun de janvier 2010 à décembre 2024, soit 180 observations. "
             "Elles sont enrichies de 14 variables exogènes macroéconomiques et climatiques : "
             "PIB, indice de production industrielle, taux d'inflation, prix du baril Brent, "
             "taux de change effectif réel (REER), dépenses publiques, solde budgétaire, "
             "indices de prix alimentaires (FAO), prix des engrais, du cacao, du café Robusta, "
             "du coton, du riz, des températures moyennes et des précipitations."),
            ("Ingénierie des variables (feature engineering)",
             "La construction des prédicteurs repose sur les lags de la série cible, "
             "la décomposition STL (tendance, saisonnalité, résidu), la sélection des "
             "décalages résiduels les plus informatifs par corrélation croisée partielle "
             "(PACF), ainsi que l'agrégation trimestrielle des variables disponibles "
             "uniquement à basse fréquence."),
            ("Familles de modèles évalués",
             "Six familles de modèles ont été comparées : modèles linéaires "
             "régularisés (Ridge, Lasso, ElasticNet) et méthodes ensemblistes ou à "
             "noyau (Random Forest, XGBoost, SVR à noyau RBF). La validation adopte "
             "un schéma de validation croisée temporelle à fenêtre expansive, évitant "
             "tout risque de fuite d'information future."),
            ("Critère de sélection du modèle optimal",
             "Le modèle retenu est celui qui présente la meilleure convergence sur "
             "quatre métriques évaluées sur l'année test 2024 : RMSE (erreur "
             "quadratique moyenne), MAE (erreur absolue moyenne), MAPE (erreur "
             "relative) et R² (coefficient de détermination). XGBoost est retenu "
             "pour les horizons H=3 et H=6, et Random Forest pour H=12."),
        ]

        for sous_titre, texte in meth_sections:
            story.append(Paragraph(sous_titre, h2_st))
            story.append(Paragraph(texte, body_st))

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    return buf.getvalue()
